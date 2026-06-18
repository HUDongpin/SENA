from pathlib import Path
from textwrap import dedent

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "exports" / "theoretical-framework"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "SENA_theoretical_framework_APA7.docx"
FIGURE_PATH = OUT_DIR / "SENA_theoretical_framework_figure.png"


TITLE = "Social-Epistemic Nexus Analytics: A Theoretical Framework for Evidence-Traceable Collaborative Learning Analytics"
SHORT_TITLE = "SENA Theoretical Framework"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:{}".format(key)), str(value))


def clear_table_borders(table):
    nil = {"val": "nil", "sz": "0", "color": "FFFFFF"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)


def set_three_line_table_borders(table):
    clear_table_borders(table)
    line = {"val": "single", "sz": "12", "color": "000000"}
    header_line = {"val": "single", "sz": "8", "color": "000000"}
    for cell in table.rows[0].cells:
        set_cell_border(cell, top=line, bottom=header_line)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=line)


def set_cell_text(cell, text, bold=False, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run()
    set_run_font(run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text="", style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, spacing=2.0):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = spacing
    if first_line and align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading_level_1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_heading_level_2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_table(doc, number, title, headers, rows, widths=None):
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_num.paragraph_format.line_spacing = 2.0
    p_num.paragraph_format.space_after = Pt(0)
    run = p_num.add_run(f"Table {number}")
    set_run_font(run, bold=True)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.line_spacing = 2.0
    p_title.paragraph_format.space_after = Pt(0)
    run = p_title.add_run(title)
    set_run_font(run, italic=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        table.rows[0].cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            set_cell_text(row.cells[idx], value)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                row.cells[idx].width = Inches(widths[idx])
    set_three_line_table_borders(table)
    doc.add_paragraph()
    return table


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(text)
    set_run_font(run)
    return p


def make_framework_figure():
    fig, ax = plt.subplots(figsize=(9.2, 5.2), dpi=220)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")

    boxes = [
        ("Learning Activity\nand Discourse Data", 0.4, 5.4, 2.4, 1.2),
        ("Validated Evidence\nUtterances, Codes,\nInteractions", 0.4, 2.9, 2.4, 1.5),
        ("Layered Constructs\nS: person-person\nW: code-code\nB/G: bridges", 3.6, 4.2, 2.5, 1.9),
        ("SENA Fusion Model\nTyped supra-adjacency\nmatrix", 6.8, 4.35, 2.3, 1.6),
        ("Interpretive Lenses\nDual Lens Dashboard\nFusion Canvas\nTemporal Trace", 9.6, 4.0, 2.1, 2.2),
        ("Educational Claims\nRoles, brokerage,\nuptake, equity,\nintervention", 6.8, 1.0, 2.3, 1.7),
        ("Human Review and\nValidity Guardrails", 9.5, 1.15, 2.2, 1.35),
    ]

    for label, x, y, w, h in boxes:
        rect = plt.Rectangle((x, y), w, h, facecolor="#F7F9FB", edgecolor="#1F4E5F", linewidth=1.4)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=8.5, family="serif")

    def arrow(x1, y1, x2, y2):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", lw=1.2, color="#1F4E5F"))

    arrow(2.8, 6.0, 3.6, 5.35)
    arrow(2.8, 3.65, 3.6, 4.65)
    arrow(6.1, 5.1, 6.8, 5.15)
    arrow(9.1, 5.15, 9.6, 5.15)
    arrow(8.0, 4.35, 8.0, 2.7)
    arrow(9.6, 4.45, 8.95, 2.7)
    arrow(9.5, 1.85, 9.1, 1.85)
    arrow(10.6, 4.0, 10.6, 2.5)

    ax.text(5.25, 7.35, "Social-Epistemic Nexus Analytics (SENA)",
            ha="center", va="center", fontsize=12, weight="bold", family="serif")
    ax.text(5.25, 7.0, "A theory of collaborative learning as socially situated epistemic work",
            ha="center", va="center", fontsize=8.5, family="serif")
    fig.tight_layout()
    fig.savefig(FIGURE_PATH, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build_document():
    make_framework_figure()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    # Title page
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(TITLE)
    set_run_font(run, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run("SENA Project Working Paper")
    set_run_font(run)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run("June 8, 2026")
    set_run_font(run)
    doc.add_page_break()

    add_heading_level_1(doc, "Abstract")
    abstract = (
        "This paper proposes a coherent theoretical framework for Social-Epistemic Nexus "
        "Analytics (SENA), a web-based research and interpretation approach that integrates "
        "Social Network Analysis, Epistemic Network Analysis, and evidence-traceable "
        "person-code bridge modeling. Building from the SENA fusion matrix and the current "
        "SENA project design, the framework conceptualizes collaborative learning as "
        "socially situated epistemic work: learners participate in relational structures, "
        "activate and connect epistemic resources, and develop roles through temporal "
        "patterns of social-epistemic contribution. The framework specifies the ontology, "
        "mechanisms, claims, analytic layers, validity safeguards, and design implications "
        "needed to make SENA theoretically coherent for educational technology and learning "
        "sciences research. It argues that SENA's contribution is not a visual overlay of "
        "SNA and ENA, but a typed, normalized, and evidence-linked representation of the "
        "social-epistemic nexus in collaborative learning."
    )
    add_paragraph(doc, abstract, first_line=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run("Keywords: ")
    set_run_font(run, italic=True)
    run = p.add_run("collaborative learning, learning analytics, social network analysis, epistemic network analysis, social-epistemic nexus, educational technology")
    set_run_font(run)
    doc.add_page_break()

    add_heading_level_1(doc, TITLE)
    add_paragraph(doc,
        "Collaborative learning is not only a matter of who interacts with whom, nor is it only a matter of which concepts appear in discourse. It is a joint process in which participation structures, discourse moves, epistemic resources, tools, tasks, norms, and time interact to produce learning opportunities. The current SENA project begins from this problem: SNA can describe social relations but is often content-blind, while ENA can describe epistemic code relations but does not by itself explain the social positions through which those relations are produced. SENA therefore requires a theory that is simultaneously educational, technological, and mathematical.")
    add_paragraph(doc,
        "The framework proposed here treats Social-Epistemic Nexus Analytics as a theory of evidence-traceable collaborative learning analytics. The term nexus is important. A signature suggests a pattern that can be detected, whereas a nexus suggests a dynamic set of relations among people, ideas, tasks, tools, norms, and time. SENA is therefore not simply an additional dashboard or a decorative fusion canvas. It is a theory of how social participation and epistemic work become mutually constitutive in learning activity.")
    add_paragraph(doc,
        "This document builds directly from the current SENA mathematical paper, the SENA web-tool development specification, the feasibility analysis, and the literature access audit. The central mathematical object is the normalized fusion matrix in which the social layer S, epistemic co-occurrence layer W, and bridge layer B are represented as a typed supra-adjacency structure. The central educational claim is that this mathematical object is meaningful only when embedded in a theory of collaborative learning, human interpretation, evidence traceability, and validity safeguards.")

    add_heading_level_1(doc, "Theoretical Commitments")
    add_paragraph(doc,
        "The first commitment is that learning is socially mediated. From a sociocultural perspective, learners do not merely acquire isolated knowledge objects; they participate in activity systems where interaction, language, tools, and norms shape what can be known and done. In SENA, this commitment motivates the social layer. Person-person ties are not treated as noise around cognition but as part of the condition under which epistemic work becomes possible.")
    add_paragraph(doc,
        "The second commitment is that learning is epistemically structured. Productive learning involves connections among evidence, explanation, questioning, critique, reflection, coordination, design, and other coded forms of disciplinary or collaborative practice. ENA gives SENA a way to represent these code-code relations without reducing discourse to simple word counts or isolated frequencies.")
    add_paragraph(doc,
        "The third commitment is that the central unit of explanation is the social-epistemic bridge. A learner may be socially central without advancing ideas, or epistemically rich without occupying a brokerage position. The educational question is often who connects which ideas, through which social relations, at which moment, and with what evidence. SENA introduces B and G to address this explanatory space.")
    add_paragraph(doc,
        "The fourth commitment is temporal. Collaborative learning develops through phases of orientation, exploration, evidence building, synthesis, reflection, and revision. A static aggregate can hide whether a group is becoming more distributed, more centralized, more evidence oriented, or more fragmented. SENA therefore requires temporal windows and trace representations as first-class theoretical objects.")
    add_paragraph(doc,
        "The fifth commitment is interpretive governance. Educational analytics should support human judgment rather than replace it. SENA's use of AI-assisted coding, automated reporting, or generated interpretations must remain tied to source evidence, human review, and explicit limitations. This is especially important because social and epistemic labels can influence teacher decisions, learner identities, and intervention priorities.")

    add_table(
        doc,
        1,
        "Core Constructs in the SENA Theoretical Framework",
        ["Construct", "Definition", "Educational Function"],
        [
            ["Social participation", "Patterns of interaction among learners, teachers, groups, or roles.", "Represents access, coordination, influence, isolation, reciprocity, and brokerage."],
            ["Epistemic connection", "Co-occurrence or structural relation among coded concepts, practices, or discourse moves.", "Represents how learners connect evidence, explanation, critique, reflection, and other epistemic resources."],
            ["Social-epistemic bridge", "A person-code or person-code-pair relation linking participation to epistemic contribution.", "Identifies who advances which ideas and how social position relates to epistemic work."],
            ["Temporal trace", "The development of S, W, B, and G across phases, turns, stanzas, or windows.", "Shows whether collaboration becomes more distributed, coherent, reflective, or fragmented over time."],
            ["Evidence provenance", "Traceable links from analytic objects back to utterances, interactions, coded segments, and model settings.", "Supports validity, transparency, interpretation, and responsible educational use."],
            ["Human review", "Researcher or teacher judgment applied to coding, modeling, interpretation, and reporting.", "Prevents automated analytics from becoming unsupported claims or high-stakes labels."]
        ],
        widths=[1.55, 2.8, 2.25],
    )

    add_heading_level_1(doc, "The SENA Ontology")
    add_paragraph(doc,
        "The SENA ontology consists of five object families: actors, epistemic resources, interaction events, contribution evidence, and interpretive claims. Actors include learners, teachers, teams, roles, or organizational units. Epistemic resources include codes such as question, hypothesis, evidence, explanation, reflection, critique, coordination, or domain-specific concepts. Interaction events include replies, mentions, co-editing, shared task work, proximity in turns, or teacher-assigned collaborations. Contribution evidence includes coded segments, utterances, confidence values, coder information, and stanza membership. Interpretive claims are statements such as a learner served as a bridge, a group developed stronger evidence-explanation links, or an intervention redistributed epistemic participation.")
    add_paragraph(doc,
        "This ontology makes an important distinction between data objects and educational claims. A matrix value is not yet a claim. A strong bridge edge indicates an observed relation between a person and a code or code pair; it does not automatically prove understanding, causality, or productive learning. The educational meaning emerges only through the alignment of data contract, coding scheme, theoretical construct, analytic model, and evidence review.")

    add_heading_level_1(doc, "The Fusion Model as a Theoretical Object")
    add_paragraph(doc,
        "The mathematical paper establishes that the SENA fusion matrix is a valid weighted typed adjacency matrix under explicit assumptions. Educationally, this means the framework can represent social and epistemic relations without pretending that people and concepts are the same kind of node. The person-person block S captures relational participation. The code-code block W captures epistemic structure. The bridge block B captures person-code contribution. The optional G block captures person-code-pair attribution, which is essential when the research question concerns who supported a specific ENA edge.")
    add_paragraph(doc,
        "The normalized fusion matrix should therefore be interpreted as a model of the learning situation, not as a universal measurement of learning quality. The layer weights alpha, beta, and gamma are theoretical commitments. A social-heavy model foregrounds participation structures; an epistemic-heavy model foregrounds idea relations; a bridge-heavy model foregrounds the coupling between people and ideas. A responsible SENA report must disclose these settings and test whether substantive interpretations are stable across plausible alternatives.")

    add_heading_level_1(doc, "Mechanisms of Social-Epistemic Learning")
    add_heading_level_2(doc, "Mechanism 1: Social Access to Epistemic Opportunity")
    add_paragraph(doc,
        "Interaction structures distribute access to ideas, feedback, authority, and participation. A learner's position in the social layer can affect which epistemic resources they encounter and whether their contributions are taken up. SENA operationalizes this mechanism through social degree, weighted strength, closeness, betweenness, reciprocity, components, and communities, but interprets those metrics through coded discourse evidence.")
    add_heading_level_2(doc, "Mechanism 2: Epistemic Coherence and Uptake")
    add_paragraph(doc,
        "Learning involves making connections among epistemic resources. In discussion, this may appear as evidence linked to explanation, critique linked to revision, or reflection linked to design decisions. SENA uses W to model these epistemic relations, while preserving the underlying coded segments that justify them.")
    add_heading_level_2(doc, "Mechanism 3: Bridge-Mediated Role Formation")
    add_paragraph(doc,
        "Collaborative roles are not reducible to participation volume. A learner may be a coordinator, evidence broker, reflective synthesizer, critique initiator, or peripheral but conceptually important contributor. SENA treats roles as emergent social-epistemic positions: patterns in S, W, B, G, and time jointly indicate role development.")
    add_heading_level_2(doc, "Mechanism 4: Temporal Reorganization")
    add_paragraph(doc,
        "Groups can become more coherent, more distributed, more centralized, or more fragmented over time. Temporal SENA trace models this development by comparing windows of activity. The educational question shifts from what the group looked like overall to how its social-epistemic organization changed during learning.")
    add_heading_level_2(doc, "Mechanism 5: Interpretive Calibration")
    add_paragraph(doc,
        "Analytics become educationally useful when teachers and researchers can calibrate claims against evidence. SENA's evidence drawer, export functions, human-reviewed interpretations, and report generator are therefore theoretical safeguards, not only interface features. They help distinguish warranted claims from attractive but unsupported visual patterns.")

    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_num.paragraph_format.line_spacing = 2.0
    r = p_num.add_run("Figure 1")
    set_run_font(r, bold=True)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.line_spacing = 2.0
    r = p_title.add_run("The SENA Social-Epistemic Nexus Framework")
    set_run_font(r, italic=True)
    doc.add_picture(str(FIGURE_PATH), width=Inches(6.4))
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_note.paragraph_format.line_spacing = 2.0
    r = p_note.add_run("Note. ")
    set_run_font(r, italic=True)
    r = p_note.add_run("The framework links learning activity data to validated evidence, layered social and epistemic constructs, the SENA fusion model, interpretive lenses, and educational claims under human review and validity guardrails.")
    set_run_font(r)

    add_heading_level_1(doc, "Theoretical Propositions")
    propositions = [
        ("Proposition 1: Content-sensitive centrality", "A learner's educational role cannot be inferred from social centrality alone; it depends on the epistemic content and evidence associated with that learner's ties."),
        ("Proposition 2: Epistemic contribution is socially situated", "Code-code connections are partly shaped by who participates, who receives uptake, and which social positions support the movement of ideas."),
        ("Proposition 3: Bridge strength predicts interpretable collaborative roles", "Learners with strong person-code and person-code-pair bridges are more likely to function as evidence brokers, synthesizers, critique initiators, or coordinators, depending on the code family involved."),
        ("Proposition 4: Alignment indicates social-epistemic coupling", "When a learner's code contribution profile aligns with the contribution profiles of their social neighbors, the group may exhibit shared epistemic orientation; low alignment may indicate productive diversity or fragmentation, depending on evidence."),
        ("Proposition 5: Temporal SENA traces reveal development", "Changes in S, W, B, and G across windows provide evidence about whether collaboration becomes more coherent, distributed, reflective, or inequitable over time."),
        ("Proposition 6: AI support must remain evidence-governed", "AI-assisted coding or interpretation can improve workflow efficiency only when outputs are linked to evidence, uncertainty, and human review.")
    ]
    for title, body in propositions:
        add_heading_level_2(doc, title)
        add_paragraph(doc, body)

    add_table(
        doc,
        2,
        "From Theoretical Claim to Analytic Test",
        ["Claim Type", "SENA Evidence", "Interpretive Test"],
        [
            ["Social participation", "Density, degree, strength, components, communities, reciprocity, and interaction evidence.", "Does the social structure create access, exclusion, brokerage, or coordination opportunities?"],
            ["Epistemic structure", "Code-code co-occurrence, ENA connection vectors, stanza windows, and representative excerpts.", "Do learners connect meaningful epistemic resources rather than merely mention many codes?"],
            ["Bridge contribution", "Person-code B and person-code-pair G with source utterances.", "Who supports which ideas or code-pair connections, and is the attribution evidence-traceable?"],
            ["Temporal development", "Windowed S, W, B, G and changes in fusion matrices.", "Does collaboration develop toward stronger synthesis, evidence use, reflection, or equitable participation?"],
            ["Group difference", "Delta fusion matrices, subgroup reports, and permutation or bootstrap checks.", "Are apparent differences stable beyond descriptive visualization?"],
            ["AI-mediated interpretation", "AI coding confidence, human review status, evidence snippets, and report limitations.", "Are generated claims transparent, reviewable, and appropriately cautious?"]
        ],
        widths=[1.65, 2.45, 2.5],
    )

    add_heading_level_1(doc, "Design Principles for SENA as Educational Technology")
    add_paragraph(doc,
        "The theoretical framework implies five design principles for SENA Fusion Studio. First, the system should begin with data validity rather than visualization. The Data Workspace must surface missing persons, unknown codes, duplicate records, unstable windows, and isolated nodes before any graph is interpreted. Second, the Model Builder should make theoretical choices visible: stanza definition, unit of analysis, normalization, alpha-beta-gamma weights, directedness, and thresholding.")
    add_paragraph(doc,
        "Third, the Fusion Canvas should support multiple interpretive modes. The explanatory layout is useful for teaching and sensemaking, but cross-layer distances should not be interpreted as statistical distances. ENA-space overlay is useful when learners or teams are ENA units, but social edges remain edge attributes rather than distance claims. Joint embedding is the only mode in which person and concept positions come from the same model, and even then the embedding method, random seed, and normalization must be reported.")
    add_paragraph(doc,
        "Fourth, the Evidence and Export workspace should treat every analytic claim as a claim requiring provenance. Clicking a social edge, concept edge, or bridge edge should reveal the interactions, stanzas, coded segments, and parameter settings that produced it. Fifth, AI should operate as a drafting and explanation partner. It may propose codes, flag anomalies, summarize patterns, or draft cautious interpretations, but human review remains the authority for educational meaning.")

    doc.add_page_break()
    add_table(
        doc,
        3,
        "Validity Threats and SENA Guardrails",
        ["Validity Threat", "Risk", "Required Guardrail"],
        [
            ["Layer-scale dominance", "One matrix overwhelms the fusion model because raw counts are not commensurable.", "Require within-layer normalization and weight sensitivity analysis."],
            ["Projection confusion", "Users interpret explanatory layouts or ENA overlays as a shared latent space.", "Display layout-specific interpretation notices and export embedding method metadata."],
            ["Coding unreliability", "Unstable codes produce unstable W, B, and G values.", "Report coding reliability, coder identity, AI confidence, and human review status."],
            ["Volume confounding", "Highly talkative learners appear epistemically important because they produce more coded segments.", "Compare raw contribution with normalized contribution, diversity, and evidence quality."],
            ["Causal overclaiming", "Observed relations are interpreted as causal effects of interaction or intervention.", "Use cautious language and require experimental, quasi-experimental, or temporal evidence for causal claims."],
            ["Automated authority", "AI-generated interpretation is treated as a conclusion rather than a draft.", "Attach evidence snippets, uncertainty notes, and explicit human-reviewed interpretation fields."]
        ],
        widths=[1.65, 2.35, 2.6],
    )

    add_heading_level_1(doc, "Research Agenda")
    add_paragraph(doc,
        "The framework supports a cumulative research agenda. Validation studies should compare SENA interpretations with expert qualitative analysis, teacher judgments, and learning outcomes. Method studies should test how normalization, stanza definitions, and layer weights affect stability. Design studies should examine whether teachers can use Fusion Canvas and Temporal Trace views to identify timely interventions. Equity studies should examine whether bridge metrics reveal whose ideas are taken up, whose contributions remain peripheral, and how task design shapes participation opportunities.")
    add_paragraph(doc,
        "A mature SENA research program should also distinguish descriptive, explanatory, predictive, and design claims. Descriptive claims summarize observed structures. Explanatory claims connect structures to theory and evidence. Predictive claims require validation against outcomes. Design claims require evidence that teachers or learners can use SENA outputs to improve participation, feedback, reflection, or knowledge building.")

    add_heading_level_1(doc, "Conclusion")
    add_paragraph(doc,
        "SENA is theoretically coherent when it is understood as a social-epistemic nexus framework rather than a technical merger of two visualizations. Its educational value lies in connecting who participates, what epistemic relations are formed, who bridges people and ideas, how these relations change over time, and what evidence warrants interpretation. Its technological value lies in making these relations inspectable, adjustable, exportable, and reviewable. Its mathematical value lies in the typed fusion matrix and its extensions. Together, these commitments make SENA a candidate framework for evidence-traceable, human-centered collaborative learning analytics.")

    doc.add_page_break()
    add_heading_level_1(doc, "References")
    references = [
        "Bowman, D. A., Swiecki, Z., Cai, Z., Eagan, B., Linder, R., Ruis, A. R., & Shaffer, D. W. (2021). The mathematical foundations of epistemic network analysis. In A. R. Ruis & S. B. Lee (Eds.), Advances in quantitative ethnography (pp. 91-105). Springer. https://doi.org/10.1007/978-3-030-67788-6_7",
        "Garrison, D. R., Anderson, T., & Archer, W. (2000). Critical inquiry in a text-based environment: Computer conferencing in higher education. The Internet and Higher Education, 2(2-3), 87-105. https://doi.org/10.1016/S1096-7516(00)00016-6",
        "Gasevic, D., Joksimovic, S., Eagan, B. R., & Shaffer, D. W. (2019). SENS: Network analytics to combine social and cognitive perspectives of collaborative learning. Computers in Human Behavior, 92, 562-577. https://doi.org/10.1016/j.chb.2018.07.003",
        "Kivela, M., Arenas, A., Barthelemy, M., Gleeson, J. P., Moreno, Y., & Porter, M. A. (2014). Multilayer networks. Journal of Complex Networks, 2(3), 203-271. https://doi.org/10.1093/comnet/cnu016",
        "Molenaar, I. (2022). Towards hybrid human-AI learning technologies. European Journal of Education, 57(4), 632-645. https://doi.org/10.1111/ejed.12527",
        "Scardamalia, M., & Bereiter, C. (2006). Knowledge building: Theory, pedagogy, and technology. In R. K. Sawyer (Ed.), The Cambridge handbook of the learning sciences (pp. 97-118). Cambridge University Press.",
        "SENA Project. (2026a). A formal analysis of the SENA fusion matrix: A typed supra-adjacency model for social-epistemic network analysis [Unpublished working paper].",
        "SENA Project. (2026b). SENA web-based tool design and development specification (Version 0.1) [Unpublished internal document].",
        "Shaffer, D. W., Collier, W., & Ruis, A. R. (2016). A tutorial on epistemic network analysis: Analyzing the structure of connections in cognitive, social, and interaction data. Journal of Learning Analytics, 3(3), 9-45. https://doi.org/10.18608/jla.2016.33.3",
        "Siebert-Evenstone, A., Arastoopour Irgens, G., Collier, W., Swiecki, Z., Ruis, A. R., & Shaffer, D. W. (2017). In search of conversational grain size: Modeling semantic structure using moving stanza windows. Journal of Learning Analytics, 4(3), 123-139. https://doi.org/10.18608/jla.2017.43.7",
        "Stahl, G., Koschmann, T., & Suthers, D. (2006). Computer-supported collaborative learning: An historical perspective. In R. K. Sawyer (Ed.), The Cambridge handbook of the learning sciences (pp. 409-426). Cambridge University Press.",
        "Sun, Y., & Han, J. (2012). Mining heterogeneous information networks: Principles and methodologies. Morgan & Claypool. https://doi.org/10.2200/S00433ED1V01Y201207DMK005",
        "Swiecki, Z., & Shaffer, D. W. (2020). iSENS: An integrated approach to combining epistemic and social network analyses. In V. Kovanovic, M. Scheffel, N. Pinkwart, & K. Verbert (Eds.), LAK20 conference proceedings: Celebrating 10 years of LAK: Shaping the future of the field (pp. 305-313). Association for Computing Machinery. https://doi.org/10.1145/3375462.3375505",
        "Vygotsky, L. S. (1978). Mind in society: The development of higher psychological processes. Harvard University Press.",
        "Wasserman, S., & Faust, K. (1994). Social network analysis: Methods and applications. Cambridge University Press.",
    ]
    for ref in references:
        add_reference(doc, ref)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)
